import csv
import datetime
import json
import os
from datetime import datetime, timedelta
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.exceptions import ObjectDoesNotExist
from django.http import HttpResponse, Http404, JsonResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.http import FileResponse
from django.utils import timezone
from django.views import View
from django.views.decorators.cache import cache_control
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_GET

from accounts.models import CustomUser
from .models import SensorData, SystemConfiguration, WaterUsage, WaterTankLevel, IrrigationEvent
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from django.utils.timezone import localtime
import pytz
import logging
import xlwt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from .sms import SMSService

logger = logging.getLogger(__name__)


def about(request):
    return render(request, 'irrigation/about.html')


def contact(request):
    return render(request, 'irrigation/contact.html')


@login_required
def help(request):
    return render(request, 'irrigation/help.html')


@login_required
def dashboard(request):
    try:
        config = SystemConfiguration.objects.get(user=request.user)
        emergency_active = config.emergency_stop
    except SystemConfiguration.DoesNotExist:
        emergency_active = False
    sensor_data = SensorData.objects.filter(user=request.user).order_by('-timestamp')[:20]

    context = {
        'emergency_active': emergency_active,
        'sensor_data': sensor_data,
    }
    return render(request, 'irrigation/dashboard.html', context)


@login_required
def download_user_manual_confirm(request):
    return render(request, 'irrigation/download_user_manual_confirm.html')


@login_required
def download_user_manual(request):
    if request.method == 'POST':
        confirm = request.POST.get('confirm', 'no')
        if confirm == 'yes':
            try:
                file_path = os.path.join(settings.BASE_DIR, 'irrigation', 'static', 'documents', 'user_guide.pdf')
                if os.path.exists(file_path):
                    response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
                    response['Content-Disposition'] = 'attachment; filename="user_guide.pdf"'
                    return response
                else:
                    raise Http404("User manual not found")
            except Exception as e:
                return render(request, 'irrigation/error.html',
                              {'error_message': 'An error occurred while downloading the manual.'})
        else:
            return redirect('dashboard')
    else:
        return redirect('download_user_manual_confirm')


def send_support_message(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        message = request.POST.get('message')

        # Construct the email subject and body
        subject = f"Support Request from {name}"
        body = f"""
        Name: {name}
        Email: {email}
        Message: {message}
        """

        # Send the email
        send_mail(
            subject=subject,
            message=body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[settings.SUPPORT_EMAIL],
            fail_silently=False,
        )

        # Notify the user that the message was sent
        messages.success(request, "Your message has been sent to support. We'll get back to you soon!")
        return redirect('help')

    return redirect('help')


@login_required
def visualize_data(request):
    """
    Render the visualization page.
    """
    return render(request, 'irrigation/visualize.html')


def privacy_policy(request):
    return render(request, 'irrigation/privacy.html')


def terms_of_service(request):
    return render(request, 'irrigation/terms.html')


@login_required
def get_sensor_data(request):
    data_type = request.GET.get('type', 'moisture')  # Only moisture data now
    user = request.user

    # Fetch the latest 50 sensor data entries for the logged-in user
    sensor_data = SensorData.objects.filter(user=user).order_by('-timestamp')[:50]

    # Prepare data for the chart (timestamps in UTC)
    labels = [data.timestamp.isoformat() + 'Z' for data in sensor_data]
    values = [getattr(data, data_type) for data in sensor_data]

    return JsonResponse({
        'labels': labels,
        'values': values
    })


@login_required
def download_data(request):
    """
    Download sensor data in CSV or Excel format with timestamps in East African Time (EAT).
    """
    format = request.GET.get('format', 'csv')

    # Fetch sensor data for the logged-in user
    sensor_data = SensorData.objects.filter(user=request.user).order_by('-timestamp')

    # Define East African Time (EAT) timezone
    eat_timezone = pytz.timezone('Africa/Nairobi')

    if format == 'csv':
        # Create CSV response
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="sensor_data.csv"'

        writer = csv.writer(response)
        writer.writerow(['Timestamp (EAT)', 'Moisture', 'Pump Status', 'Threshold'])

        for data in sensor_data:
            # Convert UTC timestamp to East African Time (EAT)
            timestamp_eat = localtime(data.timestamp, timezone=eat_timezone).strftime('%Y-%m-%d %H:%M:%S')
            writer.writerow([
                timestamp_eat,
                data.moisture,
                data.pump_status,
                data.threshold
            ])

        return response

    elif format == 'excel':
        # Create Excel response (requires openpyxl)
        from openpyxl import Workbook

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="sensor_data.xlsx"'

        wb = Workbook()
        ws = wb.active
        ws.title = "Sensor Data"

        # Add headers
        ws.append(['Timestamp (EAT)', 'Moisture', 'Pump Status', 'Threshold'])

        # Add data
        for data in sensor_data:
            # Convert UTC timestamp to East African Time (EAT)
            timestamp_eat = localtime(data.timestamp, timezone=eat_timezone).strftime('%Y-%m-%d %H:%M:%S')
            ws.append([
                timestamp_eat,
                data.moisture,
                data.pump_status,
                data.threshold
            ])

        wb.save(response)
        return response

    else:
        return HttpResponse("Invalid format", status=400)


@login_required
def control_panel(request):
    """
    View for the control panel page.
    """
    return render(request, 'irrigation/control_panel.html')


def keep_alive(request):
    return JsonResponse({"status": "OK"}, status=200)


class EnvCheckView(LoginRequiredMixin, View):
    """View to verify environment variable access"""

    def get(self, request, *args, **kwargs):
        env_vars = {
            'DEBUG': str(settings.DEBUG),
            'ENVIRONMENT': os.getenv('ENVIRONMENT', 'Not set'),
            'DB_HOST': os.getenv('DB_HOST', 'Not set'),
            'SECRET_KEY': '*****' if os.getenv('SECRET_KEY') else 'Not set',
            'EGOSMS_CONFIG': {
                'API_URL': settings.EGOSMS_CONFIG.get('API_URL', 'Not set'),
                'USERNAME': '*****' if settings.EGOSMS_CONFIG.get('USERNAME') else 'Not set',
                'TEST_MODE': str(settings.EGOSMS_CONFIG.get('TEST_MODE', 'Not set'))
            }
        }

        logger.info("Environment variables checked")
        return JsonResponse({
            'status': 'success',
            'environment_vars': env_vars
        })


@csrf_exempt
def trigger_notifications(request):
    """Endpoint for cron service to trigger SMS notifications"""
    # Authentication
    if request.headers.get('X-CRON-TOKEN') != settings.CRON_SECRET_KEY:
        logger.warning("Unauthorized cron attempt")
        return JsonResponse({'status': 'error', 'message': 'Unauthorized'}, status=401)

    # Process notifications
    try:
        latest_data = SensorData.objects.latest('timestamp')
        users = CustomUser.objects.filter(
            is_active=True,
            phone_number__isnull=False
        ).exclude(phone_number='')

        results = []
        for user in users:
            success, msg = SMSService.send_alert(user, latest_data)
            results.append({
                'user': user.username,
                'phone': user.phone_number[:4] + '*****',
                'status': 'success' if success else 'failed',
                'message': msg
            })
            logger.info(f"Processed {user.username}")

        return JsonResponse({'status': 'success', 'results': results})

    except ObjectDoesNotExist:
        logger.error("No sensor data available")
        return JsonResponse({'status': 'error', 'message': 'No sensor data'}, status=404)
    except Exception as e:
        logger.error(f"Cron job failed: {str(e)}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_GET
@cache_control(max_age=86400)
def manifest_view(request):
    """Serve the web app manifest with proper icon paths"""
    manifest_data = {
        "name": "Intelligent Irrigation System",
        "short_name": "IrrigationSystem",
        "description": "Smart irrigation management system",
        "start_url": "/",
        "display": "standalone",
        "background_color": "#ffffff",
        "theme_color": "#10b981",
        "icons": [
            {
                "src": "/static/irrigation/images/icon-72x72.png",
                "sizes": "72x72",
                "type": "image/png"
            },
            {
                "src": "/static/irrigation/images/icon-96x96.png",
                "sizes": "96x96",
                "type": "image/png"
            },
            {
                "src": "/static/irrigation/images/icon-144x144.png",
                "sizes": "144x144",
                "type": "image/png"
            },
            {
                "src": "/static/irrigation/images/icon-192x192.png",
                "sizes": "192x192",
                "type": "image/png"
            },
            {
                "src": "/static/irrigation/images/icon-512x512.png",
                "sizes": "512x512",
                "type": "image/png"
            }
        ]
    }

    response = HttpResponse(
        json.dumps(manifest_data, indent=2),
        content_type='application/manifest+json'
    )
    return response


@cache_control(max_age=86400)
def favicon_view(request):
    """Serve a simple favicon"""
    ico_data = (b'\x00\x00\x01\x00\x01\x00\x01\x01\x00\x00\x01\x00\x18\x00(\x00\x00\x00\x16\x00\x00\x00('
                b'\x00\x00\x00\x01\x00\x00\x00\x02\x00\x00\x00\x01\x00\x18\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00'
                b'\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00')

    return HttpResponse(ico_data, content_type='image/x-icon')


@login_required
def water_usage(request):
    """Render the water usage tracking page"""
    return render(request, 'irrigation/water_usage.html')


@login_required
def irrigation_frequency(request):
    """Render the irrigation frequency analysis page"""
    return render(request, 'irrigation/irrigation_frequency.html')


@login_required
def download_irrigation_report(request):
    """
    Download irrigation report in Excel or Word format
    """
    format_type = request.GET.get('format', 'excel')
    days = int(request.GET.get('days', 30))

    # Fetch data for the report
    user = request.user
    end_date = timezone.now()
    start_date = end_date - timedelta(days=days)

    # Get East African Time (EAT) timezone
    eat_timezone = pytz.timezone('Africa/Nairobi')

    # Convert current time to EAT for the report generation timestamp
    current_time_eat = timezone.now().astimezone(eat_timezone)

    # Get sensor data
    sensor_data = SensorData.objects.filter(
        user=user,
        timestamp__gte=start_date
    ).order_by('-timestamp')

    # Get irrigation events
    irrigation_events = IrrigationEvent.objects.filter(
        user=user,
        start_time__gte=start_date,
        completed=True
    ).order_by('-start_time')

    # Get water usage data
    water_usage = WaterUsage.objects.filter(
        user=user,
        timestamp__gte=start_date
    ).order_by('-timestamp')

    # Get tank level data
    tank_levels = WaterTankLevel.objects.filter(
        user=user,
        timestamp__gte=start_date
    ).order_by('-timestamp')

    # Calculate statistics
    total_irrigations = irrigation_events.count()
    total_water_used = sum(event.water_used_liters for event in irrigation_events)
    avg_irrigation_duration = sum(
        event.duration_minutes for event in irrigation_events) / total_irrigations if total_irrigations > 0 else 0
    avg_moisture_before = sum(event.moisture_before for event in irrigation_events if
                              event.moisture_before) / total_irrigations if total_irrigations > 0 else 0
    avg_moisture_after = sum(event.moisture_after for event in irrigation_events if
                             event.moisture_after) / total_irrigations if total_irrigations > 0 else 0

    # Prepare report data
    report_data = {
        'generated_date': current_time_eat.strftime('%Y-%m-%d %H:%M:%S'),
        'period_days': days,
        'period_start': start_date.strftime('%Y-%m-%d'),
        'period_end': end_date.strftime('%Y-%m-%d'),
        'user_name': user.username,
        'user_email': user.email,

        # Statistics
        'total_irrigations': total_irrigations,
        'total_water_used': round(total_water_used, 2),
        'avg_irrigation_duration': round(avg_irrigation_duration, 1),
        'avg_moisture_before': round(avg_moisture_before, 1),
        'avg_moisture_after': round(avg_moisture_after, 1),

        # Lists for detailed tables
        'irrigation_events': irrigation_events[:100],  # Limit to 100 events
        'sensor_readings': sensor_data[:100],
        'water_usage_records': water_usage[:100],
        'tank_levels': tank_levels[:100]
    }

    if format_type == 'excel':
        return generate_excel_report(report_data)
    elif format_type == 'word':
        return generate_word_report(report_data)
    else:
        return HttpResponse("Invalid format", status=400)


def generate_excel_report(data):
    """Generate Excel report with multiple sheets"""
    response = HttpResponse(content_type='application/vnd.ms-excel')
    response[
        'Content-Disposition'] = f'attachment; filename="irrigation_report_{data["period_start"]}_to_{data["period_end"]}.xls"'

    wb = xlwt.Workbook(encoding='utf-8')

    # Sheet 1: Summary
    ws_summary = wb.add_sheet('Summary')

    # Title
    title_style = xlwt.XFStyle()
    title_font = xlwt.Font()
    title_font.name = 'Arial'
    title_font.bold = True
    title_font.height = 16 * 20
    title_style.font = title_font

    ws_summary.write_merge(0, 0, 0, 3, 'Smart Irrigation System - Report', title_style)

    # Report info
    info_style = xlwt.XFStyle()
    info_font = xlwt.Font()
    info_font.name = 'Arial'
    info_font.height = 10 * 20
    info_style.font = info_font

    ws_summary.write(2, 0, 'Generated on:', info_style)
    ws_summary.write(2, 1, data['generated_date'], info_style)
    ws_summary.write(3, 0, 'Period:', info_style)
    ws_summary.write(3, 1, f"{data['period_start']} to {data['period_end']}", info_style)
    ws_summary.write(4, 0, 'User:', info_style)
    ws_summary.write(4, 1, data['user_name'], info_style)

    # Statistics
    stats_style = xlwt.XFStyle()
    stats_font = xlwt.Font()
    stats_font.name = 'Arial'
    stats_font.bold = True
    stats_font.height = 12 * 20
    stats_style.font = stats_font

    ws_summary.write(6, 0, 'Statistics', stats_style)

    stats_data = [
        ('Total Irrigations', data['total_irrigations']),
        ('Total Water Used (L)', data['total_water_used']),
        ('Average Irrigation Duration (min)', data['avg_irrigation_duration']),
        ('Average Moisture Before (%)', data['avg_moisture_before']),
        ('Average Moisture After (%)', data['avg_moisture_after']),
    ]

    for i, (label, value) in enumerate(stats_data):
        ws_summary.write(7 + i, 0, label)
        ws_summary.write(7 + i, 1, value)

    # Sheet 2: Irrigation Events
    ws_events = wb.add_sheet('Irrigation Events')
    headers = ['Date', 'Time', 'Duration (min)', 'Water Used (L)', 'Trigger', 'Moisture Before', 'Moisture After']

    for col, header in enumerate(headers):
        ws_events.write(0, col, header, stats_style)

    for row, event in enumerate(data['irrigation_events'], 1):
        eat_time = event.start_time.astimezone(pytz.timezone('Africa/Nairobi'))
        ws_events.write(row, 0, eat_time.strftime('%Y-%m-%d'))
        ws_events.write(row, 1, eat_time.strftime('%H:%M:%S'))
        ws_events.write(row, 2, round(event.duration_minutes, 1) if event.duration_minutes else 0)
        ws_events.write(row, 3, round(event.water_used_liters, 2) if event.water_used_liters else 0)
        ws_events.write(row, 4, event.get_trigger_reason_display())
        ws_events.write(row, 5, event.moisture_before if event.moisture_before else 'N/A')
        ws_events.write(row, 6, event.moisture_after if event.moisture_after else 'N/A')

    # Sheet 3: Sensor Readings
    ws_sensors = wb.add_sheet('Sensor Readings')
    headers = ['Timestamp (EAT)', 'Moisture (%)', 'Pump Status', 'Threshold (%)']

    for col, header in enumerate(headers):
        ws_sensors.write(0, col, header, stats_style)

    for row, sensor in enumerate(data['sensor_readings'], 1):
        eat_time = sensor.timestamp.astimezone(pytz.timezone('Africa/Nairobi'))
        ws_sensors.write(row, 0, eat_time.strftime('%Y-%m-%d %H:%M:%S'))
        ws_sensors.write(row, 1, sensor.moisture if sensor.moisture else 'N/A')
        ws_sensors.write(row, 2, 'ON' if sensor.pump_status else 'OFF')
        ws_sensors.write(row, 3, sensor.threshold)

    # Sheet 4: Water Usage
    ws_water = wb.add_sheet('Water Usage')
    headers = ['Timestamp (EAT)', 'Volume Used (L)', 'Initial Volume (L)', 'Final Volume (L)', 'Period (min)']

    for col, header in enumerate(headers):
        ws_water.write(0, col, header, stats_style)

    for row, usage in enumerate(data['water_usage_records'], 1):
        eat_time = usage.timestamp.astimezone(pytz.timezone('Africa/Nairobi'))
        period_minutes = usage.measurement_period.total_seconds() / 60
        ws_water.write(row, 0, eat_time.strftime('%Y-%m-%d %H:%M:%S'))
        ws_water.write(row, 1, round(usage.volume_used, 2))
        ws_water.write(row, 2, round(usage.initial_volume, 2))
        ws_water.write(row, 3, round(usage.final_volume, 2))
        ws_water.write(row, 4, round(period_minutes, 1))

    wb.save(response)
    return response


def generate_word_report(data):
    """Generate Word document report"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response[
        'Content-Disposition'] = f'attachment; filename="irrigation_report_{data["period_start"]}_to_{data["period_end"]}.docx"'

    document = Document()

    # Title
    title = document.add_heading('Smart Irrigation System - Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Report info
    document.add_paragraph(f"Generated on: {data['generated_date']}")
    document.add_paragraph(f"Period: {data['period_start']} to {data['period_end']}")
    document.add_paragraph(f"User: {data['user_name']} ({data['user_email']})")

    # Summary section
    document.add_heading('Executive Summary', level=1)

    document.add_paragraph(f"""
    This report summarizes irrigation activity for the period of {data['period_start']} to {data['period_end']}.

    During this period:
    • The system performed {data['total_irrigations']} irrigation events
    • Total water consumption was {data['total_water_used']} liters
    • Average irrigation duration was {data['avg_irrigation_duration']} minutes
    • Average soil moisture before irrigation: {data['avg_moisture_before']}%
    • Average soil moisture after irrigation: {data['avg_moisture_after']}%
    """)

    # Water usage analysis
    document.add_heading('Water Usage Analysis', level=1)

    document.add_paragraph(f"""
    Based on the analysis, here are key insights:

    Daily Water Consumption: {data['total_water_used'] / data['period_days']:.2f} liters/day
    Water per Irrigation: {data['total_water_used'] / data['total_irrigations'] if data['total_irrigations'] > 0 else 0:.2f} liters/irrigation

    {'⚠️ High water usage detected. Consider checking for leaks or adjusting irrigation schedule.' if data['total_water_used'] / data['period_days'] > 20 else '✓ Water usage is within normal range.'}
    """)

    # Irrigation frequency analysis
    document.add_heading('Irrigation Frequency Analysis', level=1)

    # Calculate intervals
    intervals = []
    events = list(data['irrigation_events'])
    for i in range(len(events) - 1):
        interval = (events[i].start_time - events[i + 1].start_time).total_seconds() / 3600
        intervals.append(abs(interval))

    if intervals:
        avg_interval = sum(intervals) / len(intervals)
        min_interval = min(intervals)
        max_interval = max(intervals)

        document.add_paragraph(f"""
        Irrigation Frequency Metrics:
        • Average interval between irrigations: {avg_interval:.1f} hours ({avg_interval / 24:.1f} days)
        • Shortest interval: {min_interval:.1f} hours
        • Longest interval: {max_interval:.1f} hours

        {'⚠️ High irrigation frequency detected. Soil may be drying too quickly. Consider adding mulch or improving soil water retention.' if avg_interval < 12 else '✓ Irrigation frequency is within normal range.'}
        """)
    else:
        document.add_paragraph("Not enough data to calculate irrigation frequency metrics.")

    # Recommendations
    document.add_heading('Recommendations', level=1)

    recommendations = []

    if data['total_water_used'] / data['period_days'] > 15:
        recommendations.append("• High water consumption detected. Review irrigation schedule and check for leaks.")

    if data['total_irrigations'] / data['period_days'] > 3:
        recommendations.append("• Frequent irrigation cycles. Consider raising moisture threshold to reduce frequency.")

    if data['avg_moisture_after'] - data['avg_moisture_before'] < 10:
        recommendations.append("• Low moisture increase per irrigation. Consider increasing irrigation duration.")

    if not recommendations:
        recommendations.append("• System performing well. Continue current settings and monitor regularly.")

    for rec in recommendations:
        document.add_paragraph(rec, style='List Bullet')

    # Irrigation Events Table
    if data['irrigation_events']:
        document.add_heading('Irrigation Events Log', level=1)

        # Add table
        table = document.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Date', 'Time', 'Duration (min)', 'Water Used (L)', 'Trigger', 'Moisture Change']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True

        # Data
        for event in data['irrigation_events'][:50]:  # Limit to 50 events
            eat_time = event.start_time.astimezone(pytz.timezone('Africa/Nairobi'))
            row_cells = table.add_row().cells

            row_cells[0].text = eat_time.strftime('%Y-%m-%d')
            row_cells[1].text = eat_time.strftime('%H:%M')
            row_cells[2].text = f"{round(event.duration_minutes, 1) if event.duration_minutes else 0}"
            row_cells[3].text = f"{round(event.water_used_liters, 2) if event.water_used_liters else 0}"
            row_cells[4].text = event.get_trigger_reason_display()

            moisture_change = ""
            if event.moisture_before and event.moisture_after:
                change = event.moisture_after - event.moisture_before
                moisture_change = f"{change:+d}% ({event.moisture_before}% → {event.moisture_after}%)"
            row_cells[5].text = moisture_change

    # Footer
    document.add_paragraph()
    document.add_paragraph(f"Report generated by Smart Irrigation System - {data['generated_date']}")

    document.save(response)
    return response

